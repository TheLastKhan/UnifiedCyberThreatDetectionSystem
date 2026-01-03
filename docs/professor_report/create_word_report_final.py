"""
CyberGuard Birleşik Proje Raporu - Final Versiyon
v2 (detaylı içerik + ekran görüntüleri) + v3 (mimari, pattern, trade-off)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

REPORT_DIR = r"c:\Users\hakan\UnifiedCyberThreatDetectionSystem\docs\professor_report"
SCREENSHOTS_DIR = os.path.join(REPORT_DIR, "screenshots")
OUTPUT_FILE = os.path.join(REPORT_DIR, "CyberGuard_Proje_Raporu_Final.docx")

def set_cell_shading(cell, color):
    """Hücre arka plan rengini ayarla"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_image(doc, filename, caption, width=6.5):
    """Görsel ve açıklama ekle"""
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        cap.paragraph_format.space_after = Pt(18)
        return True
    return False

def create_table_with_header(doc, headers, data, header_color="003366"):
    """Başlıklı tablo oluştur"""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, header_color)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True
    
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            row.cells[col_idx].text = str(cell_data)
    
    return table

def create_report():
    doc = Document()
    
    # Sayfa ayarları
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ═══════════════════════════════════════════════════════════════════
    # KAPAK SAYFASI
    # ═══════════════════════════════════════════════════════════════════
    
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CyberGuard")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Birleşik Siber Tehdit Tespit Sistemi")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    doc.add_paragraph()
    
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Teknik Dokümantasyon ve Proje Raporu")
    run.font.size = Pt(16)
    
    for _ in range(6):
        doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Versiyon: 2.0.0\n").bold = True
    info.add_run("Tarih: Ocak 2026")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # İÇİNDEKİLER
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("İÇİNDEKİLER", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc = [
        ("1. YÖNETİCİ ÖZETİ", "3"),
        ("2. SİSTEM GENEL BAKIŞ", "4"),
        ("   2.1 Amaç ve Hedefler", "4"),
        ("   2.2 Kapsam", "4"),
        ("   2.3 Teknoloji Yığını", "5"),
        ("3. YAZILIM MİMARİSİ VE TASARIM", "6"),
        ("   3.1 Mimari Karakterizasyon", "6"),
        ("   3.2 Mimari Kararların Gerekçeleri", "7"),
        ("   3.3 Katman Ayrımı ve Sorumluluklar", "8"),
        ("4. MİMARİ KALIPLAR VE TASARIM DESENLERİ", "9"),
        ("   4.1 Pattern-Mapping Tablosu", "9"),
        ("   4.2 Kalıp Seçim Gerekçeleri", "10"),
        ("5. KULLANICI ARAYÜZÜ", "11"),
        ("   5.1 Ana Panel (Dashboard)", "11"),
        ("   5.2 E-posta Analizi", "13"),
        ("   5.3 Web Log Analizi", "15"),
        ("   5.4 Korelasyon Analizi", "17"),
        ("   5.5 Model Karşılaştırma", "19"),
        ("   5.6 Raporlar ve Ayarlar", "21"),
        ("6. YAPAY ZEKA MODELLERİ", "23"),
        ("   6.1 BERT (DistilBERT)", "23"),
        ("   6.2 FastText", "24"),
        ("   6.3 TF-IDF + Random Forest", "24"),
        ("7. TEST METODOLOJİSİ VE SONUÇLARI", "25"),
        ("   7.1 Test Stratejisi ve Amacı", "25"),
        ("   7.2 Fonksiyonel Test Sonuçları", "26"),
        ("8. MODEL KARŞILAŞTIRMASI VE TRADE-OFF ANALİZİ", "27"),
        ("   8.1 Performans Karşılaştırması", "27"),
        ("   8.2 Hız vs Doğruluk Trade-off", "28"),
        ("   8.3 False Positive/Negative Analizi", "29"),
        ("   8.4 Concept Drift Riski", "30"),
        ("9. API REFERANSI", "31"),
        ("10. KURULUM VE YAPILANDIRMA", "32"),
    ]
    
    for item, page in toc:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run("\t" * 5 + page)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 1. YÖNETİCİ ÖZETİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("1. YÖNETİCİ ÖZETİ", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard, kurumsal siber güvenlik ihtiyaçlarına yönelik geliştirilmiş, "
        "yapay zeka destekli bir tehdit tespit platformudur. Sistem, e-posta tabanlı "
        "phishing saldırıları ile web tabanlı saldırıları (SQL Injection, XSS, DDoS) "
        "gerçek zamanlı olarak tespit etme kapasitesine sahiptir."
    )
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("Temel Özellikler", level=2)
    
    features_data = [
        ("Çoklu AI Modeli", "BERT, FastText ve TF-IDF olmak üzere üç farklı yapay zeka modeli"),
        ("Web Log Analizi", "Isolation Forest algoritması ile anomali tespiti"),
        ("Korelasyon Analizi", "E-posta ve web tehditlerinin zaman ve IP bazlı ilişkilendirilmesi"),
        ("Gerçek Zamanlı Dashboard", "Chart.js ile interaktif grafikler ve anlık istatistikler"),
        ("Çoklu Dil Desteği", "Türkçe ve İngilizce kullanıcı arayüzü"),
        ("Docker Deployment", "Altı container ile hazır dağıtım altyapısı"),
        ("REST API", "15+ endpoint ile tam entegrasyon imkanı"),
    ]
    
    create_table_with_header(doc, ["Özellik", "Açıklama"], features_data)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 2. SİSTEM GENEL BAKIŞ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("2. SİSTEM GENEL BAKIŞ", level=1)
    
    doc.add_heading("2.1 Amaç ve Hedefler", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Birincil Amaç: ").bold = True
    p.add_run(
        "Kurumsal ortamlarda e-posta ve web tabanlı siber tehditleri yapay zeka "
        "teknolojileri kullanarak otomatik olarak tespit etmek ve raporlamak."
    )
    p.paragraph_format.space_after = Pt(12)
    
    objectives = [
        "Phishing e-postalarını %90+ doğrulukla tespit etmek",
        "Web saldırı girişimlerini gerçek zamanlı olarak belirlemek",
        "Farklı vektörlerden gelen tehditleri ilişkilendirmek",
        "Güvenlik analistlerine kullanımı kolay bir arayüz sunmak",
        "Mevcut güvenlik altyapılarına API üzerinden entegre olmak",
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading("2.2 Kapsam", level=2)
    
    scope_data = [
        ("E-posta phishing tespiti", "Ağ trafiği analizi"),
        ("Web log anomali analizi", "Endpoint koruma"),
        ("Tehdit korelasyonu", "Malware analizi"),
        ("Raporlama ve export", "Otomatik müdahale"),
    ]
    
    create_table_with_header(doc, ["Kapsam İçi", "Kapsam Dışı"], scope_data, "006633")
    
    doc.add_heading("2.3 Teknoloji Yığını", level=2)
    
    tech_data = [
        ("Backend", "Python, Flask, Gunicorn", "3.8+, 2.0+, 21.0+"),
        ("Frontend", "HTML5, CSS3, JavaScript, Chart.js", "ES6+, 4.0+"),
        ("Veritabanı", "PostgreSQL, SQLAlchemy", "15.0, 2.0+"),
        ("Önbellek", "Redis", "7.0+"),
        ("AI/ML", "scikit-learn, PyTorch, Transformers", "1.0+, 2.0+, 4.0+"),
        ("NLP", "NLTK, spaCy, FastText", "3.8+, 3.0+, -"),
        ("Konteynerizasyon", "Docker, Docker Compose", "24.0+, 2.0+"),
        ("İzleme", "Prometheus, Grafana", "2.45+, 10.0+"),
    ]
    
    create_table_with_header(doc, ["Katman", "Teknoloji", "Versiyon"], tech_data)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 3. YAZILIM MİMARİSİ VE TASARIM (YENİ - Hocanın İstediği)
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("3. YAZILIM MİMARİSİ VE TASARIM", level=1)
    
    doc.add_heading("3.1 Mimari Karakterizasyon", level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard, modüler, servis-odaklı bir mimari üzerine inşa edilmiştir. "
        "Sistemin mimari karakteri şu şekilde tanımlanabilir:"
    )
    p.paragraph_format.space_after = Pt(12)
    
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Cm(1)
    quote.paragraph_format.right_indent = Cm(1)
    run = quote.add_run(
        '"CyberGuard is designed as a modular, service-oriented architecture where the '
        'sensing logic and presentation layers are separated, which allows machine learning '
        'models to develop independently."'
    )
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    quote.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("Mimari Tipi: Request-Response + Event-Driven Hybrid", level=3)
    
    p = doc.add_paragraph()
    p.add_run(
        "Sistem temel olarak request-response paradigmasını kullanmakla birlikte, "
        "tehdit tespiti ve korelasyon analizi bileşenlerinde event-driven yaklaşımı benimser:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    paradigm_data = [
        ("Dashboard → API", "Request-Response", "Kullanıcı istekleri synchronous olarak işlenir"),
        ("Email/Web Log → Detection", "Event-Driven", "Gelen veriler event olarak işlenir"),
        ("Detection → Correlation", "Publisher-Subscriber", "Tehditler korelasyon motoruna publish edilir"),
        ("Correlation → Alerts", "Event-Driven", "Koordineli saldırılarda alert event'leri oluşur"),
    ]
    
    create_table_with_header(doc, ["Bileşen", "Paradigma", "Açıklama"], paradigm_data)
    
    doc.add_page_break()
    
    doc.add_heading("3.2 Mimari Kararların Gerekçeleri", level=2)
    
    doc.add_heading("Neden Phishing ve Web Log Aynı Backend'de?", level=3)
    
    p = doc.add_paragraph()
    p.add_run("Karar: ").bold = True
    p.add_run("E-posta phishing tespiti ve web log analizi tek bir Flask API backend'inde birleştirilmiştir.")
    p.paragraph_format.space_after = Pt(8)
    
    p = doc.add_paragraph()
    p.add_run("Gerekçe:").bold = True
    
    reasons1 = [
        "Korelasyon Avantajı: Aynı IP adresinden gelen phishing e-postası ve web saldırısı hızlıca ilişkilendirilebilir",
        "Kaynak Verimliliği: Tek container, düşük memory footprint",
        "Deployment Basitliği: Tek docker image, kolay bakım ve güncelleme",
        "Veri Tutarlılığı: Merkezi PostgreSQL veritabanı, single source of truth",
    ]
    
    for reason in reasons1:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_heading("Neden Model Inference API İçinde?", level=3)
    
    p = doc.add_paragraph()
    p.add_run("Karar: ").bold = True
    p.add_run("ML modelleri doğrudan Flask API container'ı içinde çalıştırılmaktadır.")
    p.paragraph_format.space_after = Pt(8)
    
    reasons2 = [
        "Latency Optimizasyonu: Model → API arası network hop'u elimine edilmiştir (~5-10ms tasarruf)",
        "Session State: Modeller bir kez yüklenir ve memory'de tutulur (cold start yok)",
        "Debugging Kolaylığı: End-to-end tracing tek process'te yapılabilir",
    ]
    
    for reason in reasons2:
        doc.add_paragraph(reason, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Trade-off: ").bold = True
    p.add_run("Bu yaklaşım horizontal scaling'i zorlaştırır. Yüksek throughput senaryolarında dedicated inference server'lara geçiş önerilir.")
    
    doc.add_heading("3.3 Katman Ayrımı ve Sorumluluklar", level=2)
    
    layers_data = [
        ("Presentation (View)", "Flask Dashboard + Jinja2 + JS", "Kullanıcı etkileşimi, visualization"),
        ("Application (Controller)", "Flask REST API Routes", "Business logic, input sanitization"),
        ("Domain (Model)", "Email/Web Detectors", "ML inference, risk scoring"),
        ("Data (Persistence)", "PostgreSQL + Redis", "Data persistence, caching"),
    ]
    
    create_table_with_header(doc, ["Katman", "Teknoloji", "Sorumluluk"], layers_data)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 4. MİMARİ KALIPLAR VE TASARIM DESENLERİ (YENİ - Pattern Mapping)
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("4. MİMARİ KALIPLAR VE TASARIM DESENLERİ", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard sistemi, bilinen birçok mimari ve tasarım modelini örtük olarak benimser. "
        "Modüler yapısı doğal olarak MVC ve olay odaklı prensiplerle uyumludur. Bu yaklaşım, "
        "sistemin "
    )
    run = p.add_run("bakım kolaylığını, ölçeklenebilirliğini ve genişletilebilirliğini ")
    run.bold = True
    p.add_run("artırır.")
    
    doc.add_heading("4.1 Pattern-Mapping Tablosu", level=2)
    
    pattern_data = [
        ("Model-View-Controller (MVC)", "Dashboard (View), Flask API (Controller), PostgreSQL + ML Models (Model)"),
        ("Event-Driven / Pub-Sub", "Email/Web log ingestion → Detection → Correlation → Alert"),
        ("Ensemble Learning", "BERT, FastText, TF-IDF weighted voting (0.5, 0.3, 0.2)"),
        ("Cache-Aside Pattern", "Redis ile dashboard istatistiklerinin cachelenmesi (TTL: 60s)"),
        ("Repository Pattern", "SQLAlchemy ORM ile database abstraction"),
        ("Factory Pattern", "get_bert_detector(), get_fasttext_detector() singleton instance'lar"),
        ("Strategy Pattern", "Tüm detectorlar predict() ve predict_with_explanation() implement eder"),
        ("Façade Pattern", "/api/email/analyze/hybrid 3 modeli tek interface'te birleştirir"),
        ("Circuit Breaker", "VirusTotal API erişilemezse ML-based detection ile devam"),
    ]
    
    create_table_with_header(doc, ["Mimari Kalıp / Tasarım Deseni", "CyberGuard'daki Karşılığı"], pattern_data)
    
    doc.add_page_break()
    
    doc.add_heading("4.2 Kalıp Seçim Gerekçeleri", level=2)
    
    doc.add_heading("Neden MVC?", level=3)
    for item in ["Separation of concerns: Frontend geliştiricisi API'yi bilmeden UI değiştirebilir",
                 "Testability: Controller logic unit test edilebilir",
                 "Reusability: Aynı API farklı frontend'lerden kullanılabilir"]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Neden Ensemble Learning?", level=3)
    for item in ["Single point of failure yok: Bir model başarısız olsa diğerleri çalışır",
                 "Accuracy boost: Ensemble genellikle tek modelden daha iyi performans",
                 "Explainability: Hangi modelin nasıl karar verdiği görülebilir"]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Neden Cache-Aside?", level=3)
    for item in ["Dashboard yükleme hızı: ~1s → ~200ms improvement",
                 "Database load reduction: Sık sorgular cache'ten karşılanır"]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 5. KULLANICI ARAYÜZÜ (v2'den - Ekran Görüntüleri ile)
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("5. KULLANICI ARAYÜZÜ", level=1)
    
    doc.add_heading("5.1 Ana Panel (Dashboard)", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("Sistemin genel durumunu ve tehdit istatistiklerini tek bakışta görüntülemek.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "01_dashboard.png", "Şekil 5.1: Ana Panel (Dashboard) Görünümü")
    
    doc.add_heading("Dashboard Bileşenleri", level=3)
    
    dash_components = [
        ("E-posta Analizi Kartı", "Sol üst", "Toplam analiz edilen e-posta ve phishing oranı"),
        ("Web Anomali Kartı", "Orta üst", "Web log analiz sayısı ve anomali oranı"),
        ("Toplam Tehdit Kartı", "Sağ üst", "Tüm vektörlerden tespit edilen tehdit sayısı"),
        ("Tehdit Dağılımı", "Sol alt", "Donut chart: Phishing vs Legitimate"),
        ("Model Performans", "Sağ alt", "Bar chart: Model bazlı doğruluk"),
    ]
    
    create_table_with_header(doc, ["Bileşen", "Konum", "İşlev"], dash_components)
    
    doc.add_page_break()
    
    doc.add_heading("5.2 E-posta Analizi", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("E-posta içeriklerini üç farklı yapay zeka modeli ile analiz ederek phishing tespiti yapmak.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "02_email_analysis.png", "Şekil 5.2: E-posta Analizi Sayfası")
    
    doc.add_heading("Giriş Alanları", level=3)
    
    email_inputs = [
        ("Email Subject", "E-postanın konu satırı. Phishing e-postaları genellikle aciliyet içeren konular kullanır."),
        ("From Address", "Gönderen e-posta adresi. Şüpheli domain'ler tespit edilir."),
        ("Email Body", "E-postanın tam metin içeriği. Ana analiz bu alan üzerinde yapılır."),
    ]
    
    for field, desc in email_inputs:
        p = doc.add_paragraph()
        p.add_run(f"• {field}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Analiz Sonuç Bölümü", level=3)
    
    p = doc.add_paragraph()
    p.add_run("Her üç model için ayrı ayrı sonuçlar gösterilir:")
    
    for model, desc in [("BERT Panel", "En yüksek doğruluklu model. Bağlamsal anlam çıkarımı yapar."),
                        ("FastText Panel", "En hızlı model. Yüksek hacimli işlemler için idealdir."),
                        ("TF-IDF Panel", "Baseline model. Açıklanabilir sonuçlar sunar (LIME).")]:
        p = doc.add_paragraph()
        p.add_run(f"• {model}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    doc.add_heading("5.3 Web Log Analizi", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("Web sunucu loglarını analiz ederek SQL Injection, XSS ve DDoS saldırılarını tespit etmek.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "03_web_analysis.png", "Şekil 5.3: Web Log Analizi Sayfası")
    
    doc.add_heading("Giriş Alanları", level=3)
    
    web_inputs = [
        ("IP Address", "İstemci IP adresi. Bilinen kötü niyetli IP'ler işaretlenir."),
        ("HTTP Method", "GET, POST, PUT, DELETE vb. Anomali tespitinde kullanılır."),
        ("Request Path", "İstenen URL yolu. SQL injection kalıpları aranır."),
        ("Status Code", "HTTP yanıt kodu. Çok sayıda 401/403 şüphelidir."),
        ("User Agent", "Tarayıcı/bot bilgisi. Otomatik araçlar tespit edilir."),
    ]
    
    for field, desc in web_inputs:
        p = doc.add_paragraph()
        p.add_run(f"• {field}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Analiz Algoritması", level=3)
    
    p = doc.add_paragraph()
    p.add_run("Kullanılan Model: ").bold = True
    p.add_run("Isolation Forest algoritması. Anomali tespiti için optimize edilmiştir.")
    
    doc.add_page_break()
    
    doc.add_heading("5.4 Korelasyon Analizi", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("E-posta ve web tehditlerini zaman ve IP bazında ilişkilendirerek koordineli saldırıları tespit etmek.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "04_correlation_analysis.png", "Şekil 5.4: Korelasyon Analizi Sayfası")
    
    doc.add_heading("Korelasyon Metrikleri", level=3)
    
    corr_metrics = [
        ("Korelasyon Skoru", "Pearson korelasyon katsayısı (-1 ile +1 arası)"),
        ("Korelasyon Gücü", "Very Weak / Weak / Moderate / Strong sınıflandırma"),
        ("Koordineli Saldırı", "Aynı saat diliminde hem e-posta hem web tehdidi"),
        ("IP Boost", "Aynı IP'den hem phishing hem web saldırısı geldiğinde bonus skor"),
    ]
    
    for metric, desc in corr_metrics:
        p = doc.add_paragraph()
        p.add_run(f"• {metric}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    doc.add_heading("5.5 Model Karşılaştırma", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("Tüm yapay zeka modellerinin performans metriklerini karşılaştırmalı olarak görüntülemek.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "05_model_comparison.png", "Şekil 5.5: Model Karşılaştırma Sayfası")
    
    doc.add_page_break()
    
    doc.add_heading("5.6 Raporlar ve Ayarlar", level=2)
    
    add_image(doc, "06_reports.png", "Şekil 5.6: Raporlar Sayfası")
    
    doc.add_heading("Dışa Aktarma (Export)", level=3)
    for opt, desc in [("Export to Excel", "Tüm tehdit verilerini .xlsx formatında indirir"),
                      ("Export to JSON", "API entegrasyonu için JSON formatında dışa aktarır")]:
        p = doc.add_paragraph()
        p.add_run(f"• {opt}: ").bold = True
        p.add_run(desc)
    
    add_image(doc, "07_settings.png", "Şekil 5.7: Ayarlar Sayfası")
    
    doc.add_heading("Ayar Seçenekleri", level=3)
    for setting, desc in [("Dark Mode", "Karanlık/Aydınlık tema tercihi, kalıcı olarak kaydedilir"),
                          ("Language", "Türkçe ve İngilizce dil desteği"),
                          ("Detection Threshold", "Phishing tespit eşiği (0.0 - 1.0)")]:
        p = doc.add_paragraph()
        p.add_run(f"• {setting}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 6. YAPAY ZEKA MODELLERİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("6. YAPAY ZEKA MODELLERİ", level=1)
    
    doc.add_heading("6.1 BERT (DistilBERT)", level=2)
    
    for key, val in [("Mimari", "Transformer tabanlı, bidirectional encoder"),
                     ("Kaynak", "Hugging Face Transformers kütüphanesi"),
                     ("Eğitim Verisi", "31,000+ e-posta (CEAS, Enron, Nigerian Fraud, SpamAssassin)"),
                     ("Doğruluk", "%94-97"),
                     ("İşlem Süresi", "~45ms / e-posta"),
                     ("Avantajı", "Bağlamsal anlam çıkarımı, kelime ilişkilerini anlama"),
                     ("Dezavantajı", "Diğer modellere göre daha yavaş")]:
        p = doc.add_paragraph()
        p.add_run(f"• {key}: ").bold = True
        p.add_run(val)
    
    doc.add_heading("6.2 FastText", level=2)
    
    for key, val in [("Mimari", "Word embedding + Linear classifier"),
                     ("Kaynak", "Facebook Research"),
                     ("Model Boyutu", "881 MB"),
                     ("Doğruluk", "%90-94"),
                     ("İşlem Süresi", "<1ms / e-posta"),
                     ("Avantajı", "Çok hızlı, büyük hacimler için ideal")]:
        p = doc.add_paragraph()
        p.add_run(f"• {key}: ").bold = True
        p.add_run(val)
    
    doc.add_heading("6.3 TF-IDF + Random Forest", level=2)
    
    for key, val in [("Mimari", "TF-IDF vektörizasyon + Random Forest ensemble"),
                     ("Doğruluk", "%89.75"),
                     ("ROC-AUC", "%97.50"),
                     ("İşlem Süresi", "~25ms / e-posta"),
                     ("Avantajı", "Açıklanabilir sonuçlar, özellik önem sıralaması (LIME)")]:
        p = doc.add_paragraph()
        p.add_run(f"• {key}: ").bold = True
        p.add_run(val)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 7. TEST METODOLOJİSİ VE SONUÇLARI (YENİ - Hocanın İstediği)
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("7. TEST METODOLOJİSİ VE SONUÇLARI", level=1)
    
    doc.add_heading("7.1 Test Stratejisi ve Amacı", level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard için tasarlanan test stratejisi, sistemin temel güvenlik fonksiyonlarının "
        "doğruluğunu ve kullanıcı deneyimini öncelikli olarak hedeflemiştir."
    )
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("Test Odak Alanları", level=3)
    
    test_focus_data = [
        ("Accuracy Testi", "ML modellerinin phishing/legitimate ayrımını doğru yapması", "🔴 Kritik"),
        ("Functional Testi", "Tüm UI bileşenlerinin ve API endpoint'lerinin çalışması", "🔴 Kritik"),
        ("Integration Testi", "Backend-Database-Cache entegrasyonu", "🟡 Yüksek"),
        ("Usability Testi", "Tema, dil, ayar kalıcılığı", "🟢 Orta"),
    ]
    
    create_table_with_header(doc, ["Test Tipi", "Amaç", "Öncelik"], test_focus_data)
    
    doc.add_heading("Neden Accuracy Ölçüldü?", level=3)
    p = doc.add_paragraph()
    p.add_run("ML-based siber güvenlik sistemlerinde False Positive ve False Negative oranları kritik öneme sahiptir:")
    for item in ["False Negative (kaçırılan phishing): Güvenlik açığı, potansiyel data breach",
                 "False Positive (yanlış alarm): Operasyonel verimlilik kaybı"]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Neden Latency Detaylı Ölçülmedi?", level=3)
    for item in ["Kullanım Senaryosu: CyberGuard on-demand analiz sistemidir, real-time stream processing değil",
                 "Acceptable Threshold: 1-2 saniye response time kullanıcı deneyimi için kabul edilebilir",
                 "Gelecek Çalışma: Production deployment'ta P95/P99 latency Grafana ile monitör edilmeli"]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Neden Load Test Yapılmadı?", level=3)
    for item in ["Hedef Kitle: Orta ölçekli kurumlar (10-100 concurrent user)",
                 "Current Capacity: Flask + Gunicorn (4 worker) bu senaryoyu karşılamaktadır",
                 "Gelecek Çalışma: Kurumsal deployment öncesi Apache JMeter ile load test yapılmalı"]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading("7.2 Fonksiyonel Test Sonuçları", level=2)
    
    test_results_data = [
        ("Dashboard yükleme ve grafikler", "✅ BAŞARILI"),
        ("E-posta phishing tespiti (3 model)", "✅ BAŞARILI"),
        ("E-posta legitimate sınıflandırma", "✅ BAŞARILI"),
        ("Web log anomali tespiti", "✅ BAŞARILI"),
        ("Korelasyon analizi hesaplama", "✅ BAŞARILI"),
        ("Koordineli saldırı tespiti", "✅ BAŞARILI"),
        ("Tema değiştirme ve kalıcılık", "✅ BAŞARILI"),
        ("Dil değiştirme (TR/EN)", "✅ BAŞARILI"),
    ]
    
    create_table_with_header(doc, ["Test", "Sonuç"], test_results_data)
    
    doc.add_paragraph()
    
    doc.add_heading("Performans Metrikleri", level=3)
    
    perf_data = [
        ("API ortalama yanıt süresi", "~200ms"),
        ("BERT analiz süresi", "~45ms"),
        ("FastText analiz süresi", "<1ms"),
        ("TF-IDF analiz süresi", "~25ms"),
        ("Dashboard tam yükleme", "<1 saniye"),
    ]
    
    create_table_with_header(doc, ["Metrik", "Ölçüm"], perf_data, "006633")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 8. MODEL KARŞILAŞTIRMASI VE TRADE-OFF ANALİZİ (YENİ)
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("8. MODEL KARŞILAŞTIRMASI VE TRADE-OFF ANALİZİ", level=1)
    
    doc.add_heading("8.1 Performans Karşılaştırması", level=2)
    
    model_perf_data = [
        ("BERT (DistilBERT)", "%94-97", "%95", "%93", "%94", "~45ms"),
        ("FastText", "%90-94", "%92", "%90", "%91", "<1ms"),
        ("TF-IDF + Random Forest", "%89.75", "%90", "%88", "%89", "~25ms"),
    ]
    
    create_table_with_header(doc, ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "Inference"], model_perf_data)
    
    doc.add_heading("Neden BERT Diğerlerinden Daha İyi Performans Gösterdi?", level=3)
    
    for item in ["Contextual Understanding: BERT kelimelerin bağlamını anlar",
                 "Transfer Learning: 1.5 milyar kelime üzerinde pre-train edilmiş",
                 "Subword Tokenization: 'PayPaI' gibi typosquatting saldırılarını yakalar",
                 "Attention Mechanism: Önemli kelimeleri ('urgent', 'verify') öğrenir"]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("8.2 Hız vs Doğruluk Trade-off", level=2)
    
    usecase_data = [
        ("Real-time Email Gateway", "FastText", "Yüksek throughput, <1ms latency"),
        ("Kritik Güvenlik Analizi", "BERT", "Accuracy kritik, latency kabul edilebilir"),
        ("Balanced / Genel Kullanım", "TF-IDF", "İyi denge, açıklanabilirlik"),
        ("Ensemble (Production)", "Üçü birlikte", "En yüksek accuracy, weighted voting"),
    ]
    
    create_table_with_header(doc, ["Senaryo", "Önerilen Model", "Gerekçe"], usecase_data)
    
    doc.add_page_break()
    
    doc.add_heading("8.3 False Positive / False Negative Analizi", level=2)
    
    doc.add_heading("False Positive Senaryoları (Meşru → Phishing)", level=3)
    for item in ["Agresif Marketing E-postaları: 'Limited time offer!', 'Act now!'",
                 "IT Departmanı Uyarıları: 'Your password will expire'",
                 "Kısa Mesajlar: Çok kısa mesajlarda model güvensiz olabiliyordu (v2.0'da düzeltildi)"]:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Mitigation: ").bold = True
    p.add_run("Whitelist domain desteği, threshold ayarı, human-in-the-loop review")
    
    doc.add_heading("False Negative Senaryoları (Phishing → Meşru)", level=3)
    for item in ["Hedefli Spear Phishing: Kişiselleştirilmiş saldırılar",
                 "Zero-Day Phishing: Training data'da olmayan yeni kampanyalar",
                 "Homograph Saldırıları: 'pаypal.com' (Kiril 'а' karakteri)"]:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Mitigation: ").bold = True
    p.add_run("VirusTotal API, domain age check, sürekli model retraining")
    
    doc.add_heading("8.4 Concept Drift Riski", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Concept Drift: ").bold = True
    p.add_run("Phishing saldırıları sürekli evrilir. 2025'te etkili olan pattern'ler 2026'da değişmiş olabilir.")
    
    p = doc.add_paragraph()
    p.add_run("Önerilen Stratejiler:").bold = True
    
    for item in ["Periyodik Retraining: Her 3-6 ayda bir model güncellemesi",
                 "Active Learning: False positive/negative feedback'lerden öğrenme",
                 "Ensemble Diversification: Farklı feature'lara dayanan modeller",
                 "Continuous Monitoring: Accuracy düşüşü için alerting"]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 9. API REFERANSI
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("9. API REFERANSI", level=1)
    
    api_data = [
        ("/api/health", "GET", "Sistem sağlık kontrolü"),
        ("/api/models/status", "GET", "Model yükleme durumları"),
        ("/api/email/analyze", "POST", "TF-IDF ile e-posta analizi"),
        ("/api/email/analyze/bert", "POST", "BERT ile e-posta analizi"),
        ("/api/email/analyze/fasttext", "POST", "FastText ile e-posta analizi"),
        ("/api/email/analyze/hybrid", "POST", "Tüm modeller ile analiz (Ensemble)"),
        ("/api/predict/web", "POST", "Web log anomali analizi"),
        ("/api/correlation/analyze", "GET", "Korelasyon analizi"),
        ("/api/dashboard/stats", "GET", "Dashboard istatistikleri"),
        ("/api/reports/export/excel", "GET", "Excel dışa aktarma"),
        ("/api/reports/export/json", "GET", "JSON dışa aktarma"),
        ("/api/settings", "GET/POST", "Ayarları getir/kaydet"),
        ("/api/demo/generate", "POST", "Demo veri oluştur"),
        ("/api/database/clear", "POST", "Verileri temizle"),
    ]
    
    create_table_with_header(doc, ["Endpoint", "Method", "Açıklama"], api_data)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 10. KURULUM VE YAPILANDIRMA
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("10. KURULUM VE YAPILANDIRMA", level=1)
    
    doc.add_heading("10.1 Hızlı Başlangıç (Docker)", level=2)
    
    docker_code = """# 1. Projeyi klonlayın
git clone https://github.com/TheLastKhan/UnifiedCyberThreatDetectionSystem.git
cd UnifiedCyberThreatDetectionSystem

# 2. Docker container'ları başlatın
docker-compose up -d

# 3. Servislere erişin
# Dashboard: http://localhost:5000
# Grafana: http://localhost:3000 (admin/admin)
# Prometheus: http://localhost:9090"""
    
    code_p = doc.add_paragraph()
    run = code_p.add_run(docker_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading("10.2 Docker Container Yapısı", level=2)
    
    docker_data = [
        ("threat-detection-api", "5000", "Flask API + ML Modelleri"),
        ("threat-detection-db", "5432", "PostgreSQL Veritabanı"),
        ("threat-detection-cache", "6379", "Redis Cache"),
        ("threat-detection-nginx", "80, 443", "Reverse Proxy"),
        ("threat-detection-prometheus", "9090", "Metrik Toplama"),
        ("threat-detection-grafana", "3000", "Görselleştirme Dashboard"),
    ]
    
    create_table_with_header(doc, ["Container", "Port", "İşlev"], docker_data)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # SONUÇ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("SONUÇ", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard, modern yapay zeka teknolojilerini kullanarak kapsamlı bir siber güvenlik "
        "çözümü sunmaktadır. Sistemin temel başarıları:"
    )
    p.paragraph_format.space_after = Pt(12)
    
    for item in ["✅ 3 farklı ML modeli ile yüksek doğrulukta phishing tespiti",
                 "✅ Modüler, servis-odaklı mimari ile bakım kolaylığı",
                 "✅ Bilinen tasarım kalıpları (MVC, Event-Driven, Ensemble) ile sağlam altyapı",
                 "✅ Gerçek zamanlı korelasyon analizi ile koordineli saldırı tespiti",
                 "✅ Trade-off bilinci ile kullanım senaryosuna uygun model seçimi",
                 "✅ Docker ile kolay dağıtım ve production-ready altyapı"]:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        "Sistem, özellikle orta ölçekli kurumlar için optimize edilmiş olup, "
        "gerektiğinde horizontal scaling ile genişletilebilir yapıdadır."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("© 2025-2026 CyberGuard Project Team")
    
    # Kaydet
    doc.save(OUTPUT_FILE)
    print(f"✅ Birleşik rapor oluşturuldu: {OUTPUT_FILE}")
    return OUTPUT_FILE

if __name__ == "__main__":
    create_report()
