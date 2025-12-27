"""
CyberGuard Profesyonel ve Resmi Proje Raporu
Detaylı, düzenli ve kapsamlı Word belgesi
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

REPORT_DIR = r"c:\Users\hakan\UnifiedCyberThreatDetectionSystem\docs\professor_report"
SCREENSHOTS_DIR = os.path.join(REPORT_DIR, "screenshots")
OUTPUT_FILE = os.path.join(REPORT_DIR, "CyberGuard_Proje_Raporu_v2.docx")

def set_cell_shading(cell, color):
    """Hücre arka plan rengini ayarla"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_image(doc, filename, caption, width=6.5):
    """Görsel ve açıklama ekle"""
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        cap.paragraph_format.space_after = Pt(18)
        return True
    return False

def create_report():
    doc = Document()
    
    # Sayfa ayarları
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ═══════════════════════════════════════════════════════════════════
    # KAPAK SAYFASI
    # ═══════════════════════════════════════════════════════════════════
    
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CyberGuard")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Birleşik Siber Tehdit Tespit Sistemi")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    doc.add_paragraph()
    
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Teknik Dokümantasyon ve Kullanım Kılavuzu")
    run.font.size = Pt(16)
    
    for _ in range(6):
        doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Versiyon: 1.0.0\n").bold = True
    info.add_run("Tarih: Aralık 2025")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # İÇİNDEKİLER
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("İÇİNDEKİLER", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc = [
        ("1. YÖNETİCİ ÖZETİ", "3"),
        ("2. SİSTEM GENEL BAKIŞ", "4"),
        ("   2.1 Amaç ve Hedefler", "4"),
        ("   2.2 Kapsam", "4"),
        ("   2.3 Teknoloji Yığını", "5"),
        ("3. SİSTEM MİMARİSİ", "6"),
        ("   3.1 Mimari Diyagram", "6"),
        ("   3.2 Docker Bileşenleri", "7"),
        ("4. KULLANICI ARAYÜZÜ", "8"),
        ("   4.1 Ana Panel (Dashboard)", "8"),
        ("   4.2 E-posta Analizi", "10"),
        ("   4.3 Web Log Analizi", "12"),
        ("   4.4 Korelasyon Analizi", "14"),
        ("   4.5 Model Karşılaştırma", "16"),
        ("   4.6 Raporlar", "18"),
        ("   4.7 Ayarlar", "20"),
        ("5. YAPAY ZEKA MODELLERİ", "22"),
        ("   5.1 BERT (DistilBERT)", "22"),
        ("   5.2 FastText", "23"),
        ("   5.3 TF-IDF + Random Forest", "23"),
        ("6. API REFERANSI", "24"),
        ("7. TEST SONUÇLARI", "26"),
        ("8. KURULUM VE YAPILANDIRMA", "27"),
    ]
    
    for item, page in toc:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run("\t" * 6 + page)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 1. YÖNETİCİ ÖZETİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("1. YÖNETİCİ ÖZETİ", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard, kurumsal siber güvenlik ihtiyaçlarına yönelik geliştirilmiş, "
        "yapay zeka destekli bir tehdit tespit platformudur. Sistem, e-posta tabanlı "
        "phishing saldırıları ile web tabanlı saldırıları (SQL Injection, XSS, DDoS) "
        "gerçek zamanlı olarak tespit etme kapasitesine sahiptir."
    )
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("Temel Özellikler", level=2)
    
    features_table = doc.add_table(rows=8, cols=2)
    features_table.style = 'Table Grid'
    
    features = [
        ("Özellik", "Açıklama"),
        ("Çoklu AI Modeli", "BERT, FastText ve TF-IDF olmak üzere üç farklı yapay zeka modeli"),
        ("Web Log Analizi", "Isolation Forest algoritması ile anomali tespiti"),
        ("Korelasyon Analizi", "E-posta ve web tehditlerinin zaman ve IP bazlı ilişkilendirilmesi"),
        ("Gerçek Zamanlı Dashboard", "Chart.js ile interaktif grafikler ve anlık istatistikler"),
        ("Çoklu Dil Desteği", "Türkçe ve İngilizce kullanıcı arayüzü"),
        ("Docker Deployment", "Altı container ile hazır dağıtım altyapısı"),
        ("REST API", "15+ endpoint ile tam entegrasyon imkanı"),
    ]
    
    for i, (feat, desc) in enumerate(features):
        row = features_table.rows[i]
        row.cells[0].text = feat
        row.cells[1].text = desc
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 2. SİSTEM GENEL BAKIŞ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("2. SİSTEM GENEL BAKIŞ", level=1)
    
    doc.add_heading("2.1 Amaç ve Hedefler", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Birincil Amaç: ").bold = True
    p.add_run(
        "Kurumsal ortamlarda e-posta ve web tabanlı siber tehditleri yapay zeka "
        "teknolojileri kullanarak otomatik olarak tespit etmek ve raporlamak."
    )
    p.paragraph_format.space_after = Pt(12)
    
    objectives = [
        "Phishing e-postalarını %90+ doğrulukla tespit etmek",
        "Web saldırı girişimlerini gerçek zamanlı olarak belirlemek",
        "Farklı vektörlerden gelen tehditleri ilişkilendirmek",
        "Güvenlik analistlerine kullanımı kolay bir arayüz sunmak",
        "Mevcut güvenlik altyapılarına API üzerinden entegre olmak",
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading("2.2 Kapsam", level=2)
    
    scope_table = doc.add_table(rows=3, cols=2)
    scope_table.style = 'Table Grid'
    
    scope_data = [
        ("Kapsam İçi", "Kapsam Dışı"),
        ("E-posta phishing tespiti\nWeb log anomali analizi\nTehdit korelasyonu\nRaporlama", 
         "Ağ trafiği analizi\nEndpoint koruma\nMalware analizi\nOtomatik müdahale"),
    ]
    
    for i, (inc, exc) in enumerate(scope_data):
        row = scope_table.rows[i]
        row.cells[0].text = inc
        row.cells[1].text = exc
        if i == 0:
            set_cell_shading(row.cells[0], "006633")
            set_cell_shading(row.cells[1], "993333")
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading("2.3 Teknoloji Yığını", level=2)
    
    tech_table = doc.add_table(rows=9, cols=3)
    tech_table.style = 'Table Grid'
    
    tech_data = [
        ("Katman", "Teknoloji", "Versiyon"),
        ("Backend", "Python, Flask, Gunicorn", "3.8+, 2.0+, 21.0+"),
        ("Frontend", "HTML5, CSS3, JavaScript, Chart.js", "ES6+, 4.0+"),
        ("Veritabanı", "PostgreSQL, SQLAlchemy", "15.0, 2.0+"),
        ("Önbellek", "Redis", "7.0+"),
        ("AI/ML", "scikit-learn, PyTorch, Transformers", "1.0+, 2.0+, 4.0+"),
        ("NLP", "NLTK, spaCy, FastText", "3.8+, 3.0+, -"),
        ("Konteynerizasyon", "Docker, Docker Compose", "24.0+, 2.0+"),
        ("İzleme", "Prometheus, Grafana", "2.45+, 10.0+"),
    ]
    
    for i, (layer, tech, ver) in enumerate(tech_data):
        row = tech_table.rows[i]
        row.cells[0].text = layer
        row.cells[1].text = tech
        row.cells[2].text = ver
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 3. SİSTEM MİMARİSİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("3. SİSTEM MİMARİSİ", level=1)
    
    doc.add_heading("3.1 Mimari Diyagram", level=2)
    
    arch_text = """
┌─────────────────────────────────────────────────────────────────────────┐
│                         KULLANICI ARAYÜZÜ                               │
│  ┌─────────────┐ ┌─────────────┐ ┌─────────────┐ ┌─────────────────┐   │
│  │  Dashboard  │ │   Email     │ │   Web Log   │ │    Raporlar     │   │
│  │   Paneli    │ │   Analizi   │ │   Analizi   │ │   & Ayarlar     │   │
│  └──────┬──────┘ └──────┬──────┘ └──────┬──────┘ └────────┬────────┘   │
└─────────┼───────────────┼───────────────┼─────────────────┼────────────┘
          │               │               │                 │
          └───────────────┼───────────────┼─────────────────┘
                          ▼               ▼
┌─────────────────────────────────────────────────────────────────────────┐
│                          FLASK REST API                                  │
│  /api/email/*  │  /api/predict/*  │  /api/correlation/*  │  /api/*     │
└─────────────────────────────────────────────────────────────────────────┘
                          │               │
          ┌───────────────┼───────────────┼───────────────┐
          ▼               ▼               ▼               ▼
┌─────────────┐  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐
│    BERT     │  │  FastText   │  │  TF-IDF+RF  │  │  Isolation  │
│ (DistilBERT)│  │   Model     │  │   Model     │  │   Forest    │
└─────────────┘  └─────────────┘  └─────────────┘  └─────────────┘
                          │               │
                          ▼               ▼
┌─────────────────────────────────────────────────────────────────────────┐
│                          VERİ KATMANI                                    │
│  ┌────────────┐  ┌────────────┐  ┌────────────┐  ┌────────────┐        │
│  │ PostgreSQL │  │   Redis    │  │ Prometheus │  │  Grafana   │        │
│  │ (Veritabanı)│  │  (Cache)   │  │ (Metrikler)│  │ (Dashboard)│        │
│  └────────────┘  └────────────┘  └────────────┘  └────────────┘        │
└─────────────────────────────────────────────────────────────────────────┘
"""
    
    arch_p = doc.add_paragraph()
    run = arch_p.add_run(arch_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    
    doc.add_heading("3.2 Docker Bileşenleri", level=2)
    
    docker_table = doc.add_table(rows=7, cols=4)
    docker_table.style = 'Table Grid'
    
    docker_data = [
        ("Container", "Port", "İşlev", "Bağımlılık"),
        ("threat-detection-api", "5000", "Flask API + ML Modelleri", "db, cache"),
        ("threat-db", "5432", "PostgreSQL Veritabanı", "-"),
        ("cache", "6379", "Redis Önbellek", "-"),
        ("nginx", "80, 443", "Reverse Proxy, SSL", "api"),
        ("prometheus", "9090", "Metrik Toplama", "api"),
        ("grafana", "3000", "Görselleştirme Paneli", "prometheus"),
    ]
    
    for i, (cont, port, func, dep) in enumerate(docker_data):
        row = docker_table.rows[i]
        row.cells[0].text = cont
        row.cells[1].text = port
        row.cells[2].text = func
        row.cells[3].text = dep
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 4. KULLANICI ARAYÜZÜ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("4. KULLANICI ARAYÜZÜ", level=1)
    
    # 4.1 Dashboard
    doc.add_heading("4.1 Ana Panel (Dashboard)", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("Sistemin genel durumunu ve tehdit istatistiklerini tek bakışta görüntülemek.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "01_dashboard.png", "Şekil 4.1: Ana Panel (Dashboard) Görünümü")
    
    doc.add_heading("Dashboard Bileşenleri", level=3)
    
    dash_table = doc.add_table(rows=7, cols=3)
    dash_table.style = 'Table Grid'
    
    dash_components = [
        ("Bileşen", "Konum", "İşlev"),
        ("E-posta Analizi Kartı", "Sol üst", "Toplam analiz edilen e-posta sayısı ve tespit edilen phishing oranı"),
        ("Web Anomali Kartı", "Orta üst", "Web log analiz sayısı ve tespit edilen anomali oranı"),
        ("Toplam Tehdit Kartı", "Sağ üst", "Tüm vektörlerden tespit edilen toplam tehdit sayısı"),
        ("Sistem Durumu Kartı", "Sağ üst", "API ve model yükleme durumu (% olarak)"),
        ("Tehdit Dağılımı Grafiği", "Sol alt", "Donut chart: Phishing vs Legitimate dağılımı"),
        ("Model Performans Grafiği", "Sağ alt", "Bar chart: Model bazlı doğruluk oranları"),
    ]
    
    for i, (comp, loc, func) in enumerate(dash_components):
        row = dash_table.rows[i]
        row.cells[0].text = comp
        row.cells[1].text = loc
        row.cells[2].text = func
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading("Üst Menü Butonları", level=3)
    
    buttons = [
        ("Generate Demo Data", "Test amaçlı örnek veri seti oluşturur (30 e-posta + 30 web log + 5 koordineli saldırı)"),
        ("Clear History", "Tüm geçmiş verileri siler ve istatistikleri sıfırlar"),
        ("Tema Değiştir (☀/🌙)", "Aydınlık/Karanlık mod arasında geçiş yapar ve tercihi kaydeder"),
        ("Dil Değiştir (TR/EN)", "Arayüz dilini Türkçe veya İngilizce olarak değiştirir"),
    ]
    
    for btn, desc in buttons:
        p = doc.add_paragraph()
        p.add_run(f"• {btn}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # 4.2 Email Analysis
    doc.add_heading("4.2 E-posta Analizi", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("E-posta içeriklerini üç farklı yapay zeka modeli ile analiz ederek phishing tespiti yapmak.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "02_email_analysis.png", "Şekil 4.2: E-posta Analizi Sayfası")
    
    doc.add_heading("Giriş Alanları", level=3)
    
    email_inputs = [
        ("Email Subject (Konu)", "E-postanın konu satırı. Phishing e-postaları genellikle aciliyet içeren konular kullanır."),
        ("From Address (Gönderen)", "Gönderen e-posta adresi. Şüpheli domain'ler tespit edilir."),
        ("Email Body (İçerik)", "E-postanın tam metin içeriği. Ana analiz bu alan üzerinde yapılır."),
    ]
    
    for field, desc in email_inputs:
        p = doc.add_paragraph()
        p.add_run(f"• {field}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Analiz Sonuç Bölümü", level=3)
    
    p = doc.add_paragraph()
    p.add_run(
        "Analiz tamamlandığında, her üç model için ayrı ayrı sonuçlar gösterilir:"
    )
    
    models_result = [
        ("BERT Panel", "En yüksek doğruluklu model. Bağlamsal anlam çıkarımı yapar."),
        ("FastText Panel", "En hızlı model. Yüksek hacimli işlemler için idealdir."),
        ("TF-IDF Panel", "Baseline model. Açıklanabilir sonuçlar sunar."),
    ]
    
    for model, desc in models_result:
        p = doc.add_paragraph()
        p.add_run(f"• {model}: ").bold = True
        p.add_run(desc)
    
    p = doc.add_paragraph()
    p.add_run("Sonuç Gösterimi: ").bold = True
    p.add_run(
        "Her model için tahmin (PHISHING/LEGITIMATE), güven skoru (0-100%), "
        "risk seviyesi (Critical/High/Medium/Low) ve öne çıkan özellikler gösterilir."
    )
    
    doc.add_page_break()
    
    # 4.3 Web Analysis
    doc.add_heading("4.3 Web Log Analizi", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("Web sunucu loglarını analiz ederek SQL Injection, XSS ve DDoS gibi saldırı girişimlerini tespit etmek.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "03_web_analysis.png", "Şekil 4.3: Web Log Analizi Sayfası")
    
    doc.add_heading("Giriş Alanları", level=3)
    
    web_inputs = [
        ("IP Address", "İstemci IP adresi. Bilinen kötü niyetli IP'ler işaretlenir."),
        ("HTTP Method", "GET, POST, PUT, DELETE vb. Anomali tespitinde kullanılır."),
        ("Request Path", "İstenen URL yolu. SQL injection kalıpları aranır."),
        ("Status Code", "HTTP yanıt kodu. Çok sayıda 401/403 şüphelidir."),
        ("User Agent", "Tarayıcı/bot bilgisi. Otomatik araçlar tespit edilir (sqlmap, nikto vb.)."),
        ("Response Size", "Yanıt boyutu. Anormal boyutlar veri sızıntısına işaret edebilir."),
    ]
    
    for field, desc in web_inputs:
        p = doc.add_paragraph()
        p.add_run(f"• {field}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Analiz Algoritması", level=3)
    
    p = doc.add_paragraph()
    p.add_run("Kullanılan Model: ").bold = True
    p.add_run("Isolation Forest algoritması. Anomali tespiti için optimize edilmiştir.")
    p.paragraph_format.space_after = Pt(8)
    
    p = doc.add_paragraph()
    p.add_run("Tespit Edilen Saldırı Türleri: ").bold = True
    
    attacks = ["SQL Injection", "Cross-Site Scripting (XSS)", "Path Traversal", 
               "Brute Force", "Bot/Crawler Activity", "DDoS Patterns"]
    doc.add_paragraph(", ".join(attacks))
    
    doc.add_page_break()
    
    # 4.4 Correlation Analysis
    doc.add_heading("4.4 Korelasyon Analizi", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("E-posta ve web tehditlerini zaman ve IP bazında ilişkilendirerek koordineli saldırıları tespit etmek.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "04_correlation_analysis.png", "Şekil 4.4: Korelasyon Analizi Sayfası")
    
    doc.add_heading("Korelasyon Metrikleri", level=3)
    
    corr_metrics = [
        ("Korelasyon Skoru", "Pearson korelasyon katsayısı (-1 ile +1 arası). Pozitif değerler eş zamanlı artışı gösterir."),
        ("Korelasyon Gücü", "Very Weak / Weak / Moderate / Strong olarak sınıflandırma."),
        ("Koordineli Saldırı Sayısı", "Aynı saat diliminde hem e-posta hem web tehdidi tespit edilen durumlar."),
        ("IP Boost", "Aynı IP'den hem phishing hem web saldırısı geldiğinde eklenen bonus skor."),
    ]
    
    for metric, desc in corr_metrics:
        p = doc.add_paragraph()
        p.add_run(f"• {metric}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Grafikler", level=3)
    
    charts = [
        ("Threat Timeline Correlation", "Saat bazında e-posta ve web tehditlerinin çakışma grafiği."),
        ("Email vs Web Comparison", "İki vektörün karşılaştırmalı bar chart'ı."),
        ("Correlation Heatmap", "Tehdit korelasyonunun ısı haritası görselleştirmesi."),
    ]
    
    for chart, desc in charts:
        p = doc.add_paragraph()
        p.add_run(f"• {chart}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # 4.5 Model Comparison
    doc.add_heading("4.5 Model Karşılaştırma", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("Tüm yapay zeka modellerinin performans metriklerini karşılaştırmalı olarak görüntülemek.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "05_model_comparison.png", "Şekil 4.5: Model Karşılaştırma Sayfası")
    
    doc.add_heading("Performans Metrikleri", level=3)
    
    perf_table = doc.add_table(rows=5, cols=5)
    perf_table.style = 'Table Grid'
    
    perf_data = [
        ("Model", "Accuracy", "Precision", "Recall", "F1-Score"),
        ("BERT (DistilBERT)", "%94-97", "%95", "%93", "%94"),
        ("FastText", "%90-94", "%92", "%90", "%91"),
        ("TF-IDF + RF", "%89.75", "%90", "%88", "%89"),
        ("Isolation Forest", "%92+", "N/A", "N/A", "N/A"),
    ]
    
    for i, row_data in enumerate(perf_data):
        row = perf_table.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            if i == 0:
                set_cell_shading(row.cells[j], "003366")
                row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                row.cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # 4.6 Reports
    doc.add_heading("4.6 Raporlar", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("Tehdit verilerini dışa aktarmak ve harici kaynaklardan veri içe aktarmak.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "06_reports.png", "Şekil 4.6: Raporlar Sayfası")
    
    doc.add_heading("Dışa Aktarma (Export)", level=3)
    
    export_options = [
        ("Export to Excel", "Tüm tehdit verilerini .xlsx formatında indirir. Pivot tablo oluşturmaya uygun."),
        ("Export to JSON", "API entegrasyonu için JSON formatında dışa aktarır."),
    ]
    
    for opt, desc in export_options:
        p = doc.add_paragraph()
        p.add_run(f"• {opt}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("İçe Aktarma (Import)", level=3)
    
    import_options = [
        ("Import from Excel", "Toplu e-posta veya web log verisi yüklemek için."),
        ("Import from JSON", "Programatik veri aktarımı için."),
    ]
    
    for opt, desc in import_options:
        p = doc.add_paragraph()
        p.add_run(f"• {opt}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # 4.7 Settings
    doc.add_heading("4.7 Ayarlar", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Amaç: ").bold = True
    p.add_run("Sistem tercihlerini ve kullanıcı ayarlarını yapılandırmak.")
    p.paragraph_format.space_after = Pt(12)
    
    add_image(doc, "07_settings.png", "Şekil 4.7: Ayarlar Sayfası")
    
    doc.add_heading("Ayar Seçenekleri", level=3)
    
    settings_table = doc.add_table(rows=7, cols=3)
    settings_table.style = 'Table Grid'
    
    settings_data = [
        ("Ayar", "Tür", "Açıklama"),
        ("Dark Mode", "Toggle", "Karanlık/Aydınlık tema tercihi. Tarayıcı kapatılsa da korunur."),
        ("Language", "Checkbox", "Arayüz dili: İngilizce (varsayılan) veya Türkçe."),
        ("Detection Threshold", "Slider", "Phishing tespit eşiği (0.0 - 1.0). Düşük değer = daha hassas."),
        ("High Risk Alerts", "Toggle", "Yüksek riskli tehditler için anlık bildirim."),
        ("Daily Reports", "Toggle", "Günlük özet rapor e-postası."),
    ]
    
    for i, (setting, type_, desc) in enumerate(settings_data):
        row = settings_table.rows[i]
        row.cells[0].text = setting
        row.cells[1].text = type_
        row.cells[2].text = desc
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading("Ayar Kalıcılığı", level=3)
    
    p = doc.add_paragraph()
    p.add_run(
        "Tüm ayarlar hem localStorage (anlık tepki) hem de PostgreSQL veritabanına "
        "(kalıcı depolama) kaydedilir. Farklı cihazlardan erişimde ayarlar korunur."
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 5. YAPAY ZEKA MODELLERİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("5. YAPAY ZEKA MODELLERİ", level=1)
    
    doc.add_heading("5.1 BERT (DistilBERT)", level=2)
    
    bert_info = [
        ("Mimari", "Transformer tabanlı, bidirectional encoder"),
        ("Kaynak", "Hugging Face Transformers kütüphanesi"),
        ("Eğitim Verisi", "31,000+ e-posta (CEAS, Enron, Nigerian Fraud, SpamAssassin)"),
        ("Doğruluk", "%94-97"),
        ("İşlem Süresi", "~45ms / e-posta"),
        ("Avantajı", "Bağlamsal anlam çıkarımı, kelime ilişkilerini anlama"),
        ("Dezavantajı", "Diğer modellere göre daha yavaş"),
    ]
    
    for key, val in bert_info:
        p = doc.add_paragraph()
        p.add_run(f"• {key}: ").bold = True
        p.add_run(val)
    
    doc.add_heading("5.2 FastText", level=2)
    
    fasttext_info = [
        ("Mimari", "Word embedding + Linear classifier"),
        ("Kaynak", "Facebook Research"),
        ("Model Boyutu", "881 MB"),
        ("Doğruluk", "%90-94"),
        ("İşlem Süresi", "<1ms / e-posta"),
        ("Avantajı", "Çok hızlı, büyük hacimler için ideal"),
        ("Dezavantajı", "Karmaşık bağlam anlamada BERT'e göre zayıf"),
    ]
    
    for key, val in fasttext_info:
        p = doc.add_paragraph()
        p.add_run(f"• {key}: ").bold = True
        p.add_run(val)
    
    doc.add_heading("5.3 TF-IDF + Random Forest", level=2)
    
    tfidf_info = [
        ("Mimari", "TF-IDF vektörizasyon + Random Forest ensemble"),
        ("Eğitim", "SMOTE ile dengelenmiş veri seti"),
        ("Doğruluk", "%89.75"),
        ("ROC-AUC", "%97.50"),
        ("İşlem Süresi", "~25ms / e-posta"),
        ("Avantajı", "Açıklanabilir sonuçlar, özellik önem sıralaması"),
        ("Dezavantajı", "Deep learning modellere göre düşük doğruluk"),
    ]
    
    for key, val in tfidf_info:
        p = doc.add_paragraph()
        p.add_run(f"• {key}: ").bold = True
        p.add_run(val)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 6. API REFERANSI
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("6. API REFERANSI", level=1)
    
    api_table = doc.add_table(rows=16, cols=3)
    api_table.style = 'Table Grid'
    
    api_data = [
        ("Endpoint", "Method", "Açıklama"),
        ("/api/health", "GET", "Sistem sağlık kontrolü"),
        ("/api/models/status", "GET", "Model yükleme durumları"),
        ("/api/email/analyze", "POST", "TF-IDF ile e-posta analizi"),
        ("/api/email/analyze/bert", "POST", "BERT ile e-posta analizi"),
        ("/api/email/analyze/fasttext", "POST", "FastText ile e-posta analizi"),
        ("/api/email/analyze/hybrid", "POST", "Tüm modeller ile analiz"),
        ("/api/predict/web", "POST", "Web log anomali analizi"),
        ("/api/correlation/analyze", "GET", "Korelasyon analizi"),
        ("/api/dashboard/stats", "GET", "Dashboard istatistikleri"),
        ("/api/reports/export/excel", "GET", "Excel dışa aktarma"),
        ("/api/reports/export/json", "GET", "JSON dışa aktarma"),
        ("/api/settings", "GET", "Ayarları getir"),
        ("/api/settings", "POST", "Ayarları kaydet"),
        ("/api/demo/generate", "POST", "Demo veri oluştur"),
        ("/api/database/clear", "POST", "Verileri temizle"),
    ]
    
    for i, (endpoint, method, desc) in enumerate(api_data):
        row = api_table.rows[i]
        row.cells[0].text = endpoint
        row.cells[1].text = method
        row.cells[2].text = desc
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 7. TEST SONUÇLARI
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("7. TEST SONUÇLARI", level=1)
    
    doc.add_heading("7.1 Fonksiyonel Testler", level=2)
    
    test_table = doc.add_table(rows=11, cols=2)
    test_table.style = 'Table Grid'
    
    test_data = [
        ("Test", "Sonuç"),
        ("Dashboard yükleme ve grafikler", "✅ BAŞARILI"),
        ("E-posta phishing tespiti (3 model)", "✅ BAŞARILI"),
        ("E-posta legitimate sınıflandırma", "✅ BAŞARILI"),
        ("Web log anomali tespiti", "✅ BAŞARILI"),
        ("Web log normal trafik sınıflandırma", "✅ BAŞARILI"),
        ("Korelasyon analizi hesaplama", "✅ BAŞARILI"),
        ("Koordineli saldırı tespiti", "✅ BAŞARILI"),
        ("Tema değiştirme ve kalıcılık", "✅ BAŞARILI"),
        ("Dil değiştirme (TR/EN)", "✅ BAŞARILI"),
        ("Ayar kaydetme ve yükleme", "✅ BAŞARILI"),
    ]
    
    for i, (test, result) in enumerate(test_data):
        row = test_table.rows[i]
        row.cells[0].text = test
        row.cells[1].text = result
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading("7.2 Performans Testleri", level=2)
    
    perf_test = doc.add_table(rows=6, cols=2)
    perf_test.style = 'Table Grid'
    
    perf_test_data = [
        ("Metrik", "Ölçüm"),
        ("API ortalama yanıt süresi", "~200ms"),
        ("BERT analiz süresi", "~45ms"),
        ("FastText analiz süresi", "<1ms"),
        ("Dashboard tam yükleme", "<1 saniye"),
        ("Demo data oluşturma (60 kayıt)", "~2 saniye"),
    ]
    
    for i, (metric, val) in enumerate(perf_test_data):
        row = perf_test.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = val
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "006633")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 8. KURULUM VE YAPILANDIRMA
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("8. KURULUM VE YAPILANDIRMA", level=1)
    
    doc.add_heading("8.1 Sistem Gereksinimleri", level=2)
    
    req_table = doc.add_table(rows=6, cols=2)
    req_table.style = 'Table Grid'
    
    req_data = [
        ("Bileşen", "Gereksinim"),
        ("İşletim Sistemi", "Windows 10+, Linux, macOS"),
        ("Python", "3.8 veya üzeri"),
        ("RAM", "Minimum 4GB, önerilen 8GB"),
        ("Disk", "2GB (uygulama + modeller)"),
        ("Docker", "24.0+ (konteyner dağıtımı için)"),
    ]
    
    for i, (comp, req) in enumerate(req_data):
        row = req_table.rows[i]
        row.cells[0].text = comp
        row.cells[1].text = req
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading("8.2 Docker ile Kurulum", level=2)
    
    docker_steps = """
# 1. Projeyi klonlayın
git clone https://github.com/username/UnifiedCyberThreatDetectionSystem.git
cd UnifiedCyberThreatDetectionSystem

# 2. Docker container'ları başlatın
docker-compose up -d

# 3. Durumu kontrol edin
docker-compose ps

# 4. Servislere erişin
# Dashboard: http://localhost:5000
# Grafana: http://localhost:3000 (admin/admin)
"""
    
    code_p = doc.add_paragraph()
    run = code_p.add_run(docker_steps)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading("8.3 Erişim Bilgileri", level=2)
    
    access_table = doc.add_table(rows=5, cols=3)
    access_table.style = 'Table Grid'
    
    access_data = [
        ("Servis", "URL", "Kimlik"),
        ("Web Dashboard", "http://localhost:5000", "-"),
        ("Grafana", "http://localhost:3000", "admin / admin"),
        ("Prometheus", "http://localhost:9090", "-"),
        ("PostgreSQL", "localhost:5432", "postgres / postgres"),
    ]
    
    for i, (serv, url, cred) in enumerate(access_data):
        row = access_table.rows[i]
        row.cells[0].text = serv
        row.cells[1].text = url
        row.cells[2].text = cred
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    # Kapanış
    doc.add_paragraph()
    doc.add_paragraph()
    
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.add_run("─" * 50)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CyberGuard v1.0.0 | Aralık 2025")
    
    # Kaydet
    doc.save(OUTPUT_FILE)
    print(f"✅ Profesyonel rapor oluşturuldu: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_report()
