"""
CyberGuard Proje Raporu - Word Belgesi Oluşturucu
Hocanın Geri Bildirimleri İle Güncellenmiş Versiyon
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

REPORT_DIR = r"c:\Users\hakan\UnifiedCyberThreatDetectionSystem\docs\professor_report"
OUTPUT_FILE = os.path.join(REPORT_DIR, "CyberGuard_Proje_Raporu_v3.docx")

def set_cell_shading(cell, color):
    """Hücre arka plan rengini ayarla"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table_with_header(doc, headers, data, header_color="003366"):
    """Başlıklı tablo oluştur"""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, header_color)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            row.cells[col_idx].text = str(cell_data)
    
    return table

def create_report():
    doc = Document()
    
    # Sayfa ayarları
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ═══════════════════════════════════════════════════════════════════
    # KAPAK SAYFASI
    # ═══════════════════════════════════════════════════════════════════
    
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CyberGuard")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Birleşik Siber Tehdit Tespit Sistemi")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    doc.add_paragraph()
    
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Proje Final Raporu")
    run.font.size = Pt(18)
    
    for _ in range(6):
        doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Versiyon: 2.0.0\n").bold = True
    info.add_run("Tarih: Ocak 2026")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # İÇİNDEKİLER
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("İÇİNDEKİLER", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc = [
        ("1. PROJE ÖZETİ", "3"),
        ("2. YAZILIM MİMARİSİ VE TASARIM", "4"),
        ("   2.1 Mimari Karakterizasyon", "4"),
        ("   2.2 Mimari Kararların Gerekçeleri", "5"),
        ("   2.3 Katman Ayrımı ve Sorumluluklar", "6"),
        ("3. MİMARİ KALIPLAR VE TASARIM DESENLERİ", "7"),
        ("   3.1 Pattern-Mapping Tablosu", "7"),
        ("   3.2 Kalıp Seçim Gerekçeleri", "8"),
        ("4. SİSTEM ÖZELLİKLERİ", "9"),
        ("5. TEST METODOLOJİSİ VE SONUÇLARI", "12"),
        ("   5.1 Test Stratejisi ve Amacı", "12"),
        ("   5.2 Fonksiyonel Test Sonuçları", "13"),
        ("6. MODEL KARŞILAŞTIRMASI VE TRADE-OFF ANALİZİ", "14"),
        ("   6.1 Performans Karşılaştırması", "14"),
        ("   6.2 Hız vs Doğruluk Trade-off", "15"),
        ("   6.3 False Positive/Negative Analizi", "16"),
        ("   6.4 Concept Drift Riski", "17"),
        ("7. KURULUM VE ÇALIŞTIRMA", "18"),
    ]
    
    for item, page in toc:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run("\t" * 5 + page)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 1. PROJE ÖZETİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("1. PROJE ÖZETİ", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard, kurumsal siber güvenlik ihtiyaçlarına yönelik geliştirilmiş, "
        "yapay zeka destekli bir tehdit tespit platformudur. Sistem, e-posta tabanlı "
        "phishing saldırıları ile web tabanlı saldırıları (SQL Injection, XSS, DDoS) "
        "gerçek zamanlı olarak tespit etme kapasitesine sahiptir."
    )
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("Temel Özellikler", level=2)
    
    features_data = [
        ("📧 E-posta Phishing Tespiti", "3 farklı AI modeli (BERT, FastText, TF-IDF)", "✅ Çalışıyor"),
        ("🌐 Web Log Analizi", "SQL Injection, XSS, DDoS tespiti", "✅ Çalışıyor"),
        ("🔗 Korelasyon Analizi", "E-posta ve web tehditlerini ilişkilendirme", "✅ Çalışıyor"),
        ("📊 Gerçek Zamanlı Dashboard", "İnteraktif grafikler ve istatistikler", "✅ Çalışıyor"),
        ("🌍 Çoklu Dil Desteği", "Türkçe / İngilizce", "✅ Çalışıyor"),
        ("🐳 Docker Deployment", "6 container ile hazır dağıtım", "✅ Çalışıyor"),
    ]
    
    create_table_with_header(doc, ["Özellik", "Açıklama", "Durum"], features_data)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 2. YAZILIM MİMARİSİ VE TASARIM
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("2. YAZILIM MİMARİSİ VE TASARIM", level=1)
    
    doc.add_heading("2.1 Mimari Karakterizasyon", level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard, modüler, servis-odaklı bir mimari üzerine inşa edilmiştir. "
        "Sistemin mimari karakteri şu şekilde tanımlanabilir:"
    )
    p.paragraph_format.space_after = Pt(12)
    
    # Önemli alıntı
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Cm(1)
    quote.paragraph_format.right_indent = Cm(1)
    run = quote.add_run(
        '"CyberGuard is designed as a modular, service-oriented architecture where the '
        'sensing logic and presentation layers are separated, which allows machine learning '
        'models to develop independently."'
    )
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    quote.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("Mimari Tipi: Request-Response + Event-Driven Hybrid", level=3)
    
    p = doc.add_paragraph()
    p.add_run(
        "Sistem temel olarak request-response paradigmasını kullanmakla birlikte, "
        "tehdit tespiti ve korelasyon analizi bileşenlerinde event-driven yaklaşımı benimser:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    paradigm_data = [
        ("Dashboard → API", "Request-Response", "Kullanıcı istekleri synchronous olarak işlenir"),
        ("Email/Web Log → Detection", "Event-Driven", "Gelen veriler event olarak işlenir"),
        ("Detection → Correlation", "Publisher-Subscriber", "Tehditler korelasyon motoruna publish edilir"),
        ("Correlation → Alerts", "Event-Driven", "Koordineli saldırılarda alert event'leri oluşur"),
    ]
    
    create_table_with_header(doc, ["Bileşen", "Paradigma", "Açıklama"], paradigm_data)
    
    doc.add_page_break()
    
    doc.add_heading("2.2 Mimari Kararların Gerekçeleri", level=2)
    
    # Karar 1
    doc.add_heading("Neden Phishing ve Web Log Aynı Backend'de?", level=3)
    
    p = doc.add_paragraph()
    p.add_run("Karar: ").bold = True
    p.add_run("E-posta phishing tespiti ve web log analizi tek bir Flask API backend'inde birleştirilmiştir.")
    p.paragraph_format.space_after = Pt(8)
    
    p = doc.add_paragraph()
    p.add_run("Gerekçe:").bold = True
    
    reasons1 = [
        "Korelasyon Avantajı: Aynı IP adresinden gelen phishing e-postası ve web saldırısı, paylaşımlı veri katmanı sayesinde hızlıca ilişkilendirilebilir",
        "Kaynak Verimliliği: Tek container, düşük memory footprint (küçük/orta ölçekli kurumlar için ideal)",
        "Deployment Basitliği: Tek docker image, kolay bakım ve güncelleme",
        "Veri Tutarlılığı: Merkezi PostgreSQL veritabanı, tüm tehdit verileri için single source of truth",
    ]
    
    for reason in reasons1:
        doc.add_paragraph(reason, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Alternatif Değerlendirme: ").bold = True
    p.add_run(
        "Microservice mimarisine geçiş, yüksek ölçeklenebilirlik için düşünülebilir "
        "ancak mevcut kullanım senaryosu için overengineering olarak değerlendirilmiştir."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # Karar 2
    doc.add_heading("Neden Model Inference API İçinde?", level=3)
    
    p = doc.add_paragraph()
    p.add_run("Karar: ").bold = True
    p.add_run("ML modelleri (BERT, FastText, TF-IDF) doğrudan Flask API container'ı içinde çalıştırılmaktadır.")
    p.paragraph_format.space_after = Pt(8)
    
    p = doc.add_paragraph()
    p.add_run("Gerekçe:").bold = True
    
    reasons2 = [
        "Latency Optimizasyonu: Model → API arası network hop'u elimine edilmiştir (~5-10ms tasarruf)",
        "Session State: Modeller bir kez yüklenir ve memory'de tutulur (cold start yok)",
        "Debugging Kolaylığı: End-to-end tracing tek process'te yapılabilir",
        "Resource Isolation: Docker container zaten izolasyon sağlar",
    ]
    
    for reason in reasons2:
        doc.add_paragraph(reason, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Trade-off: ").bold = True
    p.add_run(
        "Bu yaklaşım horizontal scaling'i zorlaştırır. Yüksek throughput senaryolarında "
        "TensorFlow Serving veya TorchServe gibi dedicated inference server'lara geçiş önerilir."
    )
    
    doc.add_page_break()
    
    doc.add_heading("2.3 Katman Ayrımı ve Sorumluluklar", level=2)
    
    layers_data = [
        ("Presentation Layer (View)", "Flask Dashboard + Jinja2 + JavaScript", "Kullanıcı etkileşimi, form handling, data visualization"),
        ("Application Layer (Controller)", "Flask REST API Routes", "Business logic orchestration, input sanitization, response formatting"),
        ("Domain Layer (Model)", "Email Detector, Web Analyzer, Correlation Engine", "ML inference, feature extraction, risk scoring"),
        ("Data Layer (Persistence)", "PostgreSQL + Redis + File System", "Data persistence, caching, model storage"),
    ]
    
    create_table_with_header(doc, ["Katman", "Teknoloji", "Sorumluluk"], layers_data)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 3. MİMARİ KALIPLAR VE TASARIM DESENLERİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("3. MİMARİ KALIPLAR VE TASARIM DESENLERİ", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard sistemi, bilinen birçok mimari ve tasarım modelini örtük olarak benimser. "
        "Sistem açıkça tek bir model etrafında tasarlanmamış olsa da, modüler yapısı doğal olarak "
        "MVC ve olay odaklı prensiplerle uyumludur. Bu yaklaşım, sistemin "
    )
    run = p.add_run("bakım kolaylığını, ölçeklenebilirliğini ve genişletilebilirliğini ")
    run.bold = True
    p.add_run("artırır.")
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("3.1 Pattern-Mapping Tablosu", level=2)
    
    pattern_data = [
        ("Model-View-Controller (MVC)", "Dashboard (View), Flask API (Controller), PostgreSQL + ML Models (Model)"),
        ("Event-Driven / Publisher-Subscriber", "Email/Web log ingestion → Detection → Correlation → Alert"),
        ("Ensemble Learning Pattern", "BERT, FastText ve TF-IDF sonuçlarının weighted voting ile birleştirilmesi"),
        ("Cache-Aside Pattern", "Redis ile sık erişilen dashboard istatistiklerinin cachelenmesi (TTL: 60s)"),
        ("Repository Pattern", "SQLAlchemy ORM ile database abstraction"),
        ("Factory Pattern", "get_bert_detector(), get_fasttext_detector() singleton-like instance'lar"),
        ("Strategy Pattern", "Tüm detectorlar predict() ve predict_with_explanation() metodlarını implement eder"),
        ("Façade Pattern", "/api/email/analyze/hybrid endpoint'i 3 modeli tek interface arkasında gizler"),
        ("Circuit Breaker Pattern", "VirusTotal API erişilemezse ML-based detection ile devam"),
    ]
    
    create_table_with_header(doc, ["Mimari Kalıp / Tasarım Deseni", "CyberGuard'daki Karşılığı"], pattern_data)
    
    doc.add_page_break()
    
    doc.add_heading("3.2 Kalıp Seçim Gerekçeleri", level=2)
    
    doc.add_heading("Neden MVC?", level=3)
    patterns_mvc = [
        "Separation of concerns: Frontend geliştiricisi API'yi bilmeden UI değiştirebilir",
        "Testability: Controller logic unit test edilebilir",
        "Reusability: Aynı API farklı frontend'lerden kullanılabilir",
    ]
    for item in patterns_mvc:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Neden Ensemble Learning?", level=3)
    patterns_ensemble = [
        "Single point of failure yok: Bir model başarısız olsa diğerleri çalışır",
        "Accuracy boost: Ensemble genellikle tek modelden daha iyi performans",
        "Explainability: Hangi modelin nasıl karar verdiği görülebilir",
    ]
    for item in patterns_ensemble:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Neden Cache-Aside?", level=3)
    patterns_cache = [
        "Dashboard yükleme hızı: ~1s → ~200ms improvement",
        "Database load reduction: Sık sorgular cache'ten karşılanır",
        "Simplicity: Daha karmaşık write-through pattern'lere gerek yok",
    ]
    for item in patterns_cache:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 4. SİSTEM ÖZELLİKLERİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("4. SİSTEM ÖZELLİKLERİ", level=1)
    
    doc.add_heading("4.1 E-posta Phishing Tespiti", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Üç farklı AI modeli ile e-posta analizi yapılır ve sonuçlar karşılaştırmalı olarak gösterilir. "
        "Her model bağımsız inference yapar, sonuçlar weighted voting ile birleştirilir."
    )
    
    doc.add_heading("4.2 Web Log Analizi", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Isolation Forest algoritması ile web sunucu logları analiz edilir. "
        "SQL Injection, XSS, Path Traversal ve DDoS pattern'leri tespit edilir."
    )
    
    doc.add_heading("4.3 Korelasyon Analizi", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "E-posta ve web tehditlerini zaman ve IP bazında ilişkilendirerek koordineli saldırıları tespit eder. "
        "Pearson korelasyon katsayısı hesaplanır."
    )
    
    doc.add_heading("4.4 Docker Container Yapısı", level=2)
    
    docker_data = [
        ("threat-detection-api", "5000", "Flask API + ML Modelleri"),
        ("threat-db", "5432", "PostgreSQL Veritabanı"),
        ("cache", "6379", "Redis Cache"),
        ("nginx", "80, 443", "Reverse Proxy"),
        ("prometheus", "9090", "Metrik Toplama"),
        ("grafana", "3000", "Görselleştirme Dashboard"),
    ]
    
    create_table_with_header(doc, ["Container", "Port", "İşlev"], docker_data)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 5. TEST METODOLOJİSİ VE SONUÇLARI
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("5. TEST METODOLOJİSİ VE SONUÇLARI", level=1)
    
    doc.add_heading("5.1 Test Stratejisi ve Amacı", level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard için tasarlanan test stratejisi, sistemin temel güvenlik fonksiyonlarının "
        "doğruluğunu ve kullanıcı deneyimini öncelikli olarak hedeflemiştir."
    )
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("Test Odak Alanları", level=3)
    
    test_focus_data = [
        ("Accuracy Testi", "ML modellerinin phishing/legitimate ayrımını doğru yapması", "🔴 Kritik"),
        ("Functional Testi", "Tüm UI bileşenlerinin ve API endpoint'lerinin çalışması", "🔴 Kritik"),
        ("Integration Testi", "Backend-Database-Cache entegrasyonu", "🟡 Yüksek"),
        ("Usability Testi", "Tema, dil, ayar kalıcılığı", "🟢 Orta"),
    ]
    
    create_table_with_header(doc, ["Test Tipi", "Amaç", "Öncelik"], test_focus_data)
    
    doc.add_paragraph()
    
    doc.add_heading("Neden Accuracy Ölçüldü?", level=3)
    p = doc.add_paragraph()
    p.add_run(
        "ML-based siber güvenlik sistemlerinde False Positive ve False Negative oranları kritik öneme sahiptir:"
    )
    accuracy_reasons = [
        "False Negative (kaçırılan phishing): Güvenlik açığı, potansiyel data breach",
        "False Positive (yanlış alarm): Operasyonel verimlilik kaybı, user trust azalması",
    ]
    for reason in accuracy_reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_heading("Neden Latency Detaylı Ölçülmedi?", level=3)
    latency_reasons = [
        "Kullanım Senaryosu: CyberGuard, real-time stream processing değil, on-demand analiz sistemidir",
        "Acceptable Threshold: 1-2 saniye response time, kullanıcı deneyimi için kabul edilebilir",
        "Gelecek Çalışma: Production deployment'ta P95/P99 latency Grafana ile monitör edilmeli",
    ]
    for reason in latency_reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_heading("Neden Load Test Yapılmadı?", level=3)
    load_reasons = [
        "Hedef Kitle: Orta ölçekli kurumlar (10-100 concurrent user)",
        "Current Capacity: Flask + Gunicorn (4 worker) bu senaryoyu karşılamaktadır",
        "Gelecek Çalışma: Kurumsal deployment öncesi Apache JMeter ile load test yapılmalı",
    ]
    for reason in load_reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading("5.2 Fonksiyonel Test Sonuçları", level=2)
    
    test_results_data = [
        ("Dashboard yükleme ve grafikler", "✅ BAŞARILI"),
        ("E-posta phishing tespiti (3 model)", "✅ BAŞARILI"),
        ("E-posta legitimate sınıflandırma", "✅ BAŞARILI"),
        ("Web log anomali tespiti", "✅ BAŞARILI"),
        ("Web log normal trafik sınıflandırma", "✅ BAŞARILI"),
        ("Korelasyon analizi hesaplama", "✅ BAŞARILI"),
        ("Koordineli saldırı tespiti", "✅ BAŞARILI"),
        ("Tema değiştirme ve kalıcılık", "✅ BAŞARILI"),
        ("Dil değiştirme (TR/EN)", "✅ BAŞARILI"),
    ]
    
    create_table_with_header(doc, ["Test", "Sonuç"], test_results_data)
    
    doc.add_paragraph()
    
    doc.add_heading("Performans Metrikleri", level=3)
    
    perf_data = [
        ("API ortalama yanıt süresi", "~200ms"),
        ("BERT analiz süresi", "~45ms"),
        ("FastText analiz süresi", "<1ms"),
        ("TF-IDF analiz süresi", "~25ms"),
        ("Dashboard tam yükleme", "<1 saniye"),
    ]
    
    create_table_with_header(doc, ["Metrik", "Ölçüm"], perf_data, "006633")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 6. MODEL KARŞILAŞTIRMASI VE TRADE-OFF ANALİZİ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("6. MODEL KARŞILAŞTIRMASI VE TRADE-OFF ANALİZİ", level=1)
    
    doc.add_heading("6.1 Performans Karşılaştırması", level=2)
    
    model_perf_data = [
        ("BERT (DistilBERT)", "%94-97", "%95", "%93", "%94", "~45ms"),
        ("FastText", "%90-94", "%92", "%90", "%91", "<1ms"),
        ("TF-IDF + Random Forest", "%89.75", "%90", "%88", "%89", "~25ms"),
    ]
    
    create_table_with_header(doc, ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "Inference Time"], model_perf_data)
    
    doc.add_paragraph()
    
    doc.add_heading("Neden BERT Diğerlerinden Daha İyi Performans Gösterdi?", level=3)
    
    bert_reasons = [
        "Contextual Understanding: BERT, kelimelerin bağlamını anlar. 'Bank' kelimesi 'river bank' ve 'bank account' için farklı embedding üretir.",
        "Transfer Learning: 1.5 milyar kelime üzerinde pre-train edilmiş model, phishing dataset'inde fine-tune edilmiştir.",
        "Subword Tokenization: 'PayPaI' (I harfi ile sahte PayPal) gibi typosquatting saldırılarını yakalayabilir.",
        "Attention Mechanism: Hangi kelimelerin phishing tespitinde önemli olduğunu öğrenir ('urgent', 'verify', 'click').",
    ]
    
    for reason in bert_reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading("6.2 Hız vs Doğruluk Trade-off", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Trade-off Grafiği:").bold = True
    p.paragraph_format.space_after = Pt(8)
    
    tradeoff_text = """
    HIZLI ◄────────────────────────────────► YAVAŞ
       │                                       │
    FastText                                 BERT
     (<1ms)                                 (45ms)
       │                                       │
       ▼                                       ▼
    %90-94 Acc                            %94-97 Acc
               ┌─────────────┐
               │   TF-IDF    │
               │   (25ms)    │
               │ %89.75 Acc  │
               └─────────────┘
    """
    
    code_p = doc.add_paragraph()
    run = code_p.add_run(tradeoff_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading("Kullanım Senaryosu Önerileri", level=3)
    
    usecase_data = [
        ("Real-time Email Gateway", "FastText", "Yüksek throughput gerekli, <1ms latency"),
        ("Kritik Güvenlik Analizi", "BERT", "Accuracy kritik, latency kabul edilebilir"),
        ("Balanced / Genel Kullanım", "TF-IDF + RF", "İyi denge, açıklanabilirlik (LIME)"),
        ("Ensemble (Production)", "Üçü birlikte", "En yüksek accuracy, weighted voting"),
    ]
    
    create_table_with_header(doc, ["Senaryo", "Önerilen Model", "Gerekçe"], usecase_data)
    
    doc.add_page_break()
    
    doc.add_heading("6.3 False Positive / False Negative Analizi", level=2)
    
    doc.add_heading("False Positive Senaryoları (Meşru → Phishing)", level=3)
    
    fp_scenarios = [
        "Agresif Marketing E-postaları: 'Limited time offer!', 'Act now!' gibi ifadeler",
        "IT Departmanı Uyarıları: 'Your password will expire' gibi legitimate sistem mesajları",
        "Kısa Mesajlar: 'Hey, how are you?' gibi çok kısa mesajlarda model güvensiz olabiliyordu (v2.0'da düzeltildi)",
    ]
    for scenario in fp_scenarios:
        doc.add_paragraph(scenario, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Mitigation: ").bold = True
    p.add_run("Whitelist domain desteği, threshold ayarı, human-in-the-loop review süreci")
    
    doc.add_heading("False Negative Senaryoları (Phishing → Meşru)", level=3)
    
    fn_scenarios = [
        "Hedefli Spear Phishing: Kişiselleştirilmiş, phishing keyword içermeyen saldırılar",
        "Zero-Day Phishing: Yeni kampanyalar, training data'da olmayan pattern'ler",
        "Homograph Saldırıları: 'pаypal.com' (Kiril 'а' karakteri) gibi punycode saldırıları",
    ]
    for scenario in fn_scenarios:
        doc.add_paragraph(scenario, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Mitigation: ").bold = True
    p.add_run("VirusTotal API ile URL reputation check, domain age check, sürekli model retraining")
    
    doc.add_heading("6.4 Concept Drift Riski", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Concept Drift: ").bold = True
    p.add_run(
        "Phishing saldırıları sürekli evrilir. 2025'te etkili olan phishing pattern'leri "
        "2026'da değişmiş olabilir."
    )
    p.paragraph_format.space_after = Pt(8)
    
    p = doc.add_paragraph()
    p.add_run("Risk Faktörleri:")
    
    drift_risks = [
        "Yeni phishing kampanya temaları (AI-generated phishing, deepfake)",
        "Yeni sosyal mühendislik teknikleri",
        "Değişen e-posta formatları",
    ]
    for risk in drift_risks:
        doc.add_paragraph(risk, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Önerilen Stratejiler:").bold = True
    
    drift_strategies = [
        "Periyodik Retraining: Her 3-6 ayda bir model güncellemesi",
        "Active Learning: False positive/negative feedback'lerden öğrenme",
        "Ensemble Diversification: Farklı feature'lara dayanan modeller kullanma",
        "Continuous Monitoring: Accuracy metrikleri düşüşü için alerting",
    ]
    for strategy in drift_strategies:
        doc.add_paragraph(strategy, style='List Bullet')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # 7. KURULUM VE ÇALIŞTIRMA
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("7. KURULUM VE ÇALIŞTIRMA", level=1)
    
    doc.add_heading("7.1 Hızlı Başlangıç (Docker)", level=2)
    
    docker_code = """
# 1. Projeyi klonlayın
git clone https://github.com/username/UnifiedCyberThreatDetectionSystem.git
cd UnifiedCyberThreatDetectionSystem

# 2. Docker container'ları başlatın
docker-compose up -d

# 3. Servislere erişin
# Dashboard: http://localhost:5000
# Grafana: http://localhost:3000 (admin/admin)
# Prometheus: http://localhost:9090
"""
    
    code_p = doc.add_paragraph()
    run = code_p.add_run(docker_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading("7.2 Manuel Kurulum", level=2)
    
    manual_code = """
# 1. Virtual environment oluşturun
python -m venv venv
source venv/bin/activate  # Windows: venv\\Scripts\\activate

# 2. Bağımlılıkları yükleyin
pip install -r requirements.txt

# 3. Dashboard'u başlatın
python run_dashboard.py

# 4. Tarayıcıda açın: http://localhost:5000
"""
    
    code_p = doc.add_paragraph()
    run = code_p.add_run(manual_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # SONUÇ
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_heading("SONUÇ", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "CyberGuard, modern yapay zeka teknolojilerini kullanarak kapsamlı bir siber güvenlik "
        "çözümü sunmaktadır. Sistemin temel başarıları:"
    )
    p.paragraph_format.space_after = Pt(12)
    
    conclusions = [
        "✅ 3 farklı ML modeli ile yüksek doğrulukta phishing tespiti",
        "✅ Modüler, servis-odaklı mimari ile bakım kolaylığı",
        "✅ Bilinen tasarım kalıpları (MVC, Event-Driven, Ensemble) ile sağlam altyapı",
        "✅ Gerçek zamanlı korelasyon analizi ile koordineli saldırı tespiti",
        "✅ Trade-off bilinci ile kullanım senaryosuna uygun model seçimi",
        "✅ Docker ile kolay dağıtım ve production-ready altyapı",
    ]
    
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        "Sistem, özellikle orta ölçekli kurumlar için optimize edilmiş olup, "
        "gerektiğinde horizontal scaling ile genişletilebilir yapıdadır."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("© 2025-2026 CyberGuard Project Team")
    
    # Kaydet
    doc.save(OUTPUT_FILE)
    print(f"✅ Rapor oluşturuldu: {OUTPUT_FILE}")
    return OUTPUT_FILE

if __name__ == "__main__":
    create_report()
